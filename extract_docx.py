from docx import Document

doc = Document('/Users/augustopadilla/Documents/GitHub/kodu-app/owner_report_sample.docx')

print("=== PARAGRAPHS ===")
for i, p in enumerate(doc.paragraphs):
    if p.text.strip():
        print(f"{i}: [{p.style.name}] {p.text.strip()}")

print("\n=== TABLES ===")
for ti, table in enumerate(doc.tables):
    print(f"\nTable {ti}:")
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        print(cells)
